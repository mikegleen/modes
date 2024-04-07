# -*- coding: utf-8 -*-
"""
    Called from stocktake.sh

    Create a report of box contents with columns for the stocktake.
    Parameters:
        1. Input CSV file.
        2. Output DOCX file.

    The input CSV file has columns:
        Serial,Current,Normal,Title,Description,Condition
"""
import argparse
from collections import defaultdict, namedtuple
import csv
import re
import sys
from docx.shared import Cm
import docx
from docx.enum.style import WD_STYLE


from utl.normalize import sphinxify, normalize_id, if_not_sphinx

DEFAULT_TEMPLATE = 'etc/templates/docx/stocktake.docx'
OUTPUT_COLUMNS = 'OK,Serial,NL,Title,Condition'.split(',')
OUTPUT_WIDTHS = '1.0,2.0,1.0,10.0,5.0'
OUTPUT_WIDTHS = [Cm(float(x)) for x in OUTPUT_WIDTHS.split(',')]
# The "title" field will contain the "description" field if the "title" is empty.
RowTuple = namedtuple('RowTuple', 'current serial normal title condition')


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def pad_loc(loc):
    """
    If the location field matches the pattern "<Letter(s)><Number(s)>" like
    "S1" then pad the number part so that it sorts correctly. Also convert the
    letter part to upper case
    :param loc: location
    :return: padded location
    """
    m = re.match(r'(\D+)(\d+)', loc)
    if m:
        g1 = m.group(1).upper()
        g2 = int(m.group(2))
        return f'{g1}{g2:03}'
    else:
        return loc


def unpad_loc(loc):
    m = re.match(r'(\D+)(\d+)', loc)
    if m:
        g1 = m.group(1).upper()
        g2 = int(m.group(2))
        return f'{g1}{g2}'
    else:
        return loc


def one_row(row):
    serial = row['Serial'].removeprefix('LDHRM.')
    loc = row['Current']
    if loc is not None:
        location = pad_loc(loc)
    else:
        location = 'unknown'
    normal = row['Normal'].removeprefix('LDHRM.')
    if normal == loc:
        normal = ''
    title = row['Title']
    description = row['Description']
    if not title:
        title = description
    condition = row['Condition']
    rowtuple = RowTuple(location, serial, normal, title, condition)
    boxdict[location].append(rowtuple)


def emit_one_box(box):
    doc.add_heading(f'Location: {unpad_loc(box)}')
    boxlist = boxdict[box]
    # The template document must contain the style.
    # Create a table then delete it. Really.
    table = doc.add_table(rows=len(boxlist) + 1, cols=len(OUTPUT_COLUMNS), style='Table Grid')
    for col, width in enumerate(OUTPUT_WIDTHS):
        for cell in table.columns[col].cells:
            cell.width = width
    cells = table.rows[0].cells
    # print(hdr_cells)
    for coln, coltitle in enumerate(OUTPUT_COLUMNS):
        # print(f'{coln=}')
        cells[coln].text = coltitle
    make_rows_bold(table.rows[0])
    for ln, rowtuple in enumerate(sorted(boxlist, key=lambda r: normalize_id(r.serial)), start=1):
        # print(f'{rowtuple=}')
        cells = table.rows[ln].cells
        # print(f'{len(cells)=}')
        cells[0].text = ''
        cells[1].text = rowtuple.serial
        cells[2].text = rowtuple.normal
        cells[3].text = rowtuple.title
        cells[4].text = rowtuple.condition
    doc.add_page_break()


def main():
    section = doc.sections[0]
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.header.paragraphs[0].text = _args.header
    for row in csvreader:
        one_row(row)
    for box in sorted(boxdict.keys()):
        emit_one_box(box)
        if _args.short:
            break
    doc.save(_args.outfile)


def getparser() -> argparse.ArgumentParser:
    """
    Called either by getargs() in this file or by Sphinx.
    :return: an argparse.ArgumentParser object
    """
    parser = argparse.ArgumentParser(description='''
        Create a report of box contents with columns for the stocktake. The
        script uses a configuration file embedded in the Python code.''')
    parser.add_argument('infile', help=sphinxify('''
        CSV file produced by xml2csv.
        ''', called_from_sphinx))
    parser.add_argument('outfile', help='''
        The output DOCX file.''')
    parser.add_argument('--header', required=True, help='''
        Set the page heading.''')
    defaultstr = f'{if_not_sphinx(DEFAULT_TEMPLATE, called_from_sphinx)}'
    parser.add_argument('-t', '--template', default=DEFAULT_TEMPLATE, help=f'''
        The template to use for the output DOCX file. {defaultstr}''')
    parser.add_argument('-s', '--short', action='store_true', help='''
        Only process one object. For debugging.''')
    parser.add_argument('-v', '--verbose', type=int, default=1, help='''
        Set the verbosity. The default is 1 which prints summary
        information.''')
    return parser


def getargs(argv):
    parser = getparser()
    args = parser.parse_args(args=argv[1:])
    return args


called_from_sphinx = True


if __name__ == '__main__':
    assert sys.version_info >= (3, 11)
    called_from_sphinx = False
    if len(sys.argv) == 1:
        sys.argv.append('-h')
    _args = getargs(sys.argv)
    csvfile = open(_args.infile)
    csvreader = csv.DictReader(csvfile)
    boxdict = defaultdict(list)
    doc = docx.Document(_args.template)
    main()
